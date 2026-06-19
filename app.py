import warnings
# Or to just mute all deprecation warnings entirely:
warnings.filterwarnings("ignore", category=DeprecationWarning)

from parser.text import extract_text
from parser.tables import extract_tables
from parser.forms import extract_forms
from renderer.docx_renderer import create_docx

#docling
from parser.docling_parser import extract_with_docling
from docx import Document
#docling


PDF_PATH = "uploads/sample.pdf"
OUTPUT_PATH = "outputs/outputB.docx"


markdown = extract_with_docling(PDF_PATH)

#docling
doc = Document()
doc.add_paragraph(markdown)

doc.save(OUTPUT_PATH)

print("Done")
#docling



#text = extract_text(PDF_PATH)
#
#tables = extract_tables(PDF_PATH)
#
#forms = extract_forms(PDF_PATH)
#
#create_docx(
#    text_pages=text,
#    tables=tables,
#    forms=forms,
#    output_path=OUTPUT_PATH
#)
#
#print("Done!")