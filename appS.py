import streamlit as st
from docling.document_converter import DocumentConverter
from docx import Document
from bs4 import BeautifulSoup
import markdown
import pandas as pd
import tempfile
import os


# ----------------------------
# DOCX HELPERS
# ----------------------------

def add_table_to_docx(doc, element):
    rows = element.find_all("tr")

    if not rows:
        return

    first_row = rows[0].find_all(["th", "td"])
    cols = len(first_row)

    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"

    for row in rows:
        cells_in_row = row.find_all(["th", "td"])
        cells = table.add_row().cells

        for i, cell in enumerate(cells_in_row):
            if i < len(cells):
                cells[i].text = cell.get_text(" ", strip=True)

    doc.add_paragraph()


def add_markdown_to_docx(doc, markdown_text):
    html = markdown.markdown(markdown_text, extensions=["tables"])
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.find_all(recursive=False):

        if not hasattr(element, "name"):
            continue

        # Headings
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            doc.add_heading(element.get_text(strip=True), level=level)

        # Paragraphs
        elif element.name == "p":
            text = element.get_text(" ", strip=True)
            if text:
                doc.add_paragraph(text)

        # Bullet lists
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style="List Bullet")

        # Numbered lists
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style="List Number")

        # Tables
        elif element.name == "table":
            add_table_to_docx(doc, element)


# ----------------------------
# CORE CONVERSION
# ----------------------------

def convert_pdf_to_docx(pdf_path):
    converter = DocumentConverter()
    result = converter.convert(pdf_path)

    doc = Document()

    doc.add_heading("Converted Document", level=1)

    # Markdown text from Docling
    markdown_text = result.document.export_to_markdown()

    add_markdown_to_docx(doc, markdown_text)

    return doc


# ----------------------------
# STREAMLIT UI
# ----------------------------

st.set_page_config(page_title="DocPingu - PDF → DOCX Converter", layout="centered")

st.image("img/DocPingu.png", caption="DocPingu")
st.title(" PDF → Word Converter ")
st.write("Upload a PDF and convert it into a clean editable Word document with headings + tables.")

uploaded_file = st.file_uploader("Upload PDF", type=["pdf"])

if uploaded_file:

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_pdf.write(uploaded_file.read())
        pdf_path = tmp_pdf.name

    st.success("PDF uploaded successfully!")

    if st.button("Convert to DOCX"):

        with st.spinner("Converting..."):
            doc = convert_pdf_to_docx(pdf_path)

            output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            doc.save(output_path)

        with open(output_path, "rb") as f:
            st.download_button(
                label="⬇️ Download DOCX",
                data=f,
                file_name="converted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success("Conversion complete!")

st.write("DocPingu © Copyright 2027. Convert Word Docs, PDF Docs and More.")