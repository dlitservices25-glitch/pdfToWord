from docling.document_converter import DocumentConverter
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import pandas as pd
import markdown
import re


INPUT_PDF = "uploads/sample.pdf"
OUTPUT_DOCX = "outputs/outputA.docx"


def add_markdown_to_docx(doc, markdown_text):
    """
    Convert markdown into structured DOCX content.
    Preserves:
    - headings
    - paragraphs
    - bullet lists
    - numbered lists
    """

    html = markdown.markdown(
        markdown_text,
        extensions=["tables"]
    )
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.children:

        if element.name is None:
            continue

        # Headings
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            doc.add_heading(element.get_text(strip=True), level=level)

        # Paragraphs
        elif element.name == "p":
            text = element.get_text(" ", strip=True)
            if text:
                doc.add_paragraph(text)

        # Bullet lists
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(
                    li.get_text(" ", strip=True),
                    style="List Bullet"
                )

        # Numbered lists
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(
                    li.get_text(" ", strip=True),
                    style="List Number"
                )

        elif element.name == "table":
            rows = element.find_all("tr")
    
            if not rows:
                continue
            
            first_row = rows[0].find_all(["th", "td"])
            cols = len(first_row)
    
            table = doc.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
    
            for row in rows:
                cells_in_row = row.find_all(["th", "td"])
                cells = table.add_row().cells
    
                for i, cell in enumerate(cells_in_row):
                    if i < len(cells):
                        cells[i].text = cell.get_text(" ", strip=True)
    
            doc.add_paragraph()

def add_dataframe_table(doc, df):
    """
    Add editable Word table from DataFrame.
    """

    rows, cols = df.shape

    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    # Header row
    for col_idx, column_name in enumerate(df.columns):
        table.cell(0, col_idx).text = str(column_name)

    # Data rows
    for row_idx in range(rows):
        for col_idx in range(cols):
            value = df.iloc[row_idx, col_idx]

            if pd.isna(value):
                value = ""

            table.cell(row_idx + 1, col_idx).text = str(value)

    doc.add_paragraph()


def main():
    print("Loading PDF with Docling...")

    converter = DocumentConverter()
    result = converter.convert(INPUT_PDF)

    print("Creating DOCX...")

    doc = Document()

    # Title
    doc.add_heading("Converted Document", level=1)

    # ---------- TEXT / STRUCTURE ----------
    markdown_text = result.document.export_to_markdown()

    add_markdown_to_docx(doc, markdown_text)

    # ---------- TABLES ----------
    if hasattr(result.document, "tables"):

        for i, table in enumerate(result.document.tables):
            try:
                doc.add_heading(f"Table {i + 1}", level=2)

                df = table.export_to_dataframe(
                    doc=result.document
                )

                add_dataframe_table(doc, df)

            except Exception as e:
                print(f"Failed table {i}: {e}")

    doc.save(OUTPUT_DOCX)

    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()