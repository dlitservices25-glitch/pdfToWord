from docx import Document

def create_docx(text_pages, tables, forms, output_path):

    doc = Document()

    # TEXT
    for page_num, page in enumerate(text_pages):

        doc.add_heading(f"Page {page_num + 1}", level=1)

        for paragraph in page:
            doc.add_paragraph(paragraph)

    # TABLES
    doc.add_page_break()
    doc.add_heading("Tables", level=1)

    for table in tables:

        if not table:
            continue

        max_cols = max(len(row) for row in table)

        rows = len(table)

        word_table = doc.add_table(rows=rows, cols=max_cols)

        for i, row in enumerate(table):

            for j, cell in enumerate(row):

                cleaned = str(cell or "").strip()

                if j < max_cols:
                    word_table.cell(i, j).text = cleaned

    # FORMS
    if forms:

        doc.add_page_break()
        doc.add_heading("Form Fields", level=1)

        for key, value in forms.items():
            doc.add_paragraph(f"{key}: {value}")

    doc.save(output_path)