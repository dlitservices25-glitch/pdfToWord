from docling.document_converter import DocumentConverter
from docx import Document
from bs4 import BeautifulSoup
import markdown
import tempfile
import os


def add_table(doc, element):
   rows = element.find_all("tr")
   if not rows:
       return

   cols = len(rows[0].find_all(["th", "td"]))

   table = doc.add_table(rows=0, cols=cols)
   table.style = "Table Grid"

   for row in rows:
       cells = row.find_all(["th", "td"])
       row_cells = table.add_row().cells

       for i, cell in enumerate(cells):
           if i < len(row_cells):
               row_cells[i].text = cell.get_text(" ", strip=True)


def convert_pdf_job(pdf_path):
   converter = DocumentConverter()
   result = converter.convert(pdf_path)

   md = result.document.export_to_markdown()
   html = markdown.markdown(md, extensions=["tables"])
   soup = BeautifulSoup(html, "html.parser")

   doc = Document()
   doc.add_heading("Converted Document", level=1)

   for el in soup.find_all(recursive=False):

       if not hasattr(el, "name"):
           continue

       if el.name.startswith("h"):
           doc.add_heading(el.get_text(strip=True), int(el.name[1]))

       elif el.name == "p":
           doc.add_paragraph(el.get_text(" ", strip=True))

       elif el.name == "ul":
           for li in el.find_all("li"):
               doc.add_paragraph(li.text, style="List Bullet")

       elif el.name == "ol":
           for li in el.find_all("li"):
               doc.add_paragraph(li.text, style="List Number")

       elif el.name == "table":
           add_table(doc, el)

   output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
   doc.save(output)

   return output
